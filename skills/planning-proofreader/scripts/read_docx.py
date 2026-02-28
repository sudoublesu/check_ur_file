#!/usr/bin/env python3
"""
read_docx.py - 从 .docx 文件提取结构化内容

Usage:
    python read_docx.py <file.docx> [output.json]

Output JSON 结构:
{
  "file": "filename.docx",
  "paragraphs": [
    {"index": 0, "text": "...", "style": "Heading 1", "level": 1}
  ],
  "tables": [
    {"index": 0, "rows": [["单元格1", "单元格2"], ...]}
  ],
  "headers": ["页眉文本"],
  "footers": ["页脚文本"]
}
"""

import sys
import json
from docx import Document
from docx.oxml.ns import qn


HEADING_STYLES = {
    'heading 1': 1, 'heading 2': 2, 'heading 3': 3,
    'heading 4': 4, 'heading 5': 5, 'heading 6': 6,
    '标题 1': 1, '标题 2': 2, '标题 3': 3,
    '标题1': 1, '标题2': 2, '标题3': 3,
}


def get_level(style_name: str) -> int:
    return HEADING_STYLES.get(style_name.lower().strip(), 0)


def extract_docx(path: str) -> dict:
    doc = Document(path)
    result = {
        "file": path,
        "paragraphs": [],
        "tables": [],
        "table_positions": {},  # {table_index: preceding_para_index_in_body}
        "headers": [],
        "footers": [],
    }

    # 段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        result["paragraphs"].append({
            "index": i,
            "text": text,
            "style": style,
            "level": get_level(style),
        })

    # 表格 + 位置追踪
    # 遍历 body 子元素，按出现顺序统计段落数和表格数，
    # 记录每张表格前一个段落（w:p）的位置索引，
    # 该索引与 add_comments.py 中 body 直接子 w:p 列表的下标一致。
    xml_para_idx = -1
    xml_tbl_idx = 0
    for child in doc.element.body:
        if child.tag == qn('w:p'):
            xml_para_idx += 1
        elif child.tag == qn('w:tbl'):
            result["table_positions"][xml_tbl_idx] = xml_para_idx
            xml_tbl_idx += 1

    for t_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        result["tables"].append({"index": t_idx, "rows": rows})

    # 页眉/页脚
    for section in doc.sections:
        if section.header:
            text = section.header.paragraphs[0].text.strip() if section.header.paragraphs else ""
            if text:
                result["headers"].append(text)
        if section.footer:
            text = section.footer.paragraphs[0].text.strip() if section.footer.paragraphs else ""
            if text:
                result["footers"].append(text)

    return result


def main():
    if len(sys.argv) < 2:
        print("Usage: python read_docx.py <file.docx> [output.json]", file=sys.stderr)
        sys.exit(1)

    path = sys.argv[1]
    data = extract_docx(path)

    if len(sys.argv) >= 3:
        with open(sys.argv[2], 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"已输出到 {sys.argv[2]}")
    else:
        print(json.dumps(data, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()
